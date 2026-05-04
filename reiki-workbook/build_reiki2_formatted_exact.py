from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.platypus import BaseDocTemplate, Frame, NextPageTemplate, PageTemplate, Paragraph, PageBreak, Spacer, Table, TableStyle


DOCX_OUT = "/Users/magicalbeing/Desktop/gohighlevel/reiki-workbook/ReikiLevel2TrainingFormatted.docx"
PDF_OUT = "/Users/magicalbeing/Desktop/gohighlevel/reiki-workbook/ReikiLevel2TrainingFormatted.pdf"

TEXT = RGBColor(47, 37, 32)
MUTED = RGBColor(85, 72, 66)
BORDER = "E7DED5"
SOFT = "FFFAF5"
WASH = "F6F0E8"
TEXT_HEX = "#2f2520"
MUTED_HEX = "#554842"
BORDER_HEX = "#e7ded5"
SOFT_HEX = "#fffaf5"
WASH_HEX = "#f6f0e8"


outline = [
    "Holy Love Experience (instilling Divine Love within)",
    "The Reiki symbols & how to use them",
    "Reiki II Placement explained & given",
    "Sending Reiki with the eyes (Gyoshi Ho)",
    "Sending Reiki with the breath (Koki Ho)",
    "Distant Healing - how to send Reiki",
    "How to conduct a complete Reiki session",
    "Hands on practice with the symbols",
    "Documenting a Reiki session",
]

sections = [
    {
        "kind": "symbol",
        "title": "THE POWER SYMBOL",
        "subtitle": "Cho Ku Rei (Cho Koo Ray)",
        "quote": "“Put the power of the universe here.”",
        "blocks": [
            ("Symbol/Kanji breakdowns:", [
                "The horizontal line represents the male energy of the universe known as “Shiva” and the vertical line represents energy coming down to earth or down the spinal column. The spiral represents feminine energy or earth energy, known as “shakti.” The spiral touches the lines 7 times representing the 7 chakras.",
                "The line represents a man asleep. The vertical line, man awake. The spiral is the man on the path of his self-actualization. The spiral is the order of creation itself. It is that which turns and directs things to their natural processes of unfoldment. ",
            ], "number"),
            ("USES:", [
                "Charges and boosts Reiki energy flow",
                "Increases the power of Reiki",
                "Can focus it on a specific area",
                "Place in the palms to empower hands",
                "Clears the room of negative energy and creates a sacred space to work ",
                "Seals the energy in; closes a session",
                "Protection for yourself, your loved ones and all things! Includes physical protection, verbal protection or psychic attack.",
                "Blessing and clearing your food and drinking water.",
                "Sending energy to an accident on the road or a traffic jam, clearing a way.",
                "Working with electrical equipment. ",
                "You can draw a large power symbol down the front of your body, spiraling at the solar plexus, then draw small symbols on each chakra. This clears your channel. ",
            ], "bullet"),
        ],
    },
    {
        "kind": "symbol",
        "title": "THE MENTAL/EMOTIONAL SYMBOL",
        "subtitle": "SEI HE KI (SAY HAY KEY)",
        "quote": "Symbol of Harmony and Balance",
        "blocks": [
            ("Symbol Breakdown:", [
                "The angles on the left side represent logical, linear left-brain. The curves on the right side represent an intuitive, creative, imaginative right brain.",
            ], "body"),
            ("", [
                "The symbol can bring up and heal those deep seated problems and behaviors of a mental or emotional nature. Heals on the level of the original cause. Brings ego and soul together. ",
            ], "body"),
            ("Uses:", [
                "Balances both sides of the brain",
                "Calms and relaxes",
                "Brings balance and healing to any emotional or mental situation such as fear, anger, worry, anxiety, depression and addiction",
                "Assists in counseling",
                "Healing relationships",
                "Helps memory (use before test taking)",
                "Enhances affirmations, causing them to enter more deeply into the subconscious mind. ",
            ], "bullet"),
            ("Seiheki Chiryo - healing bad habits", [
                "This technique can be used to heal bad habits, change or release an unwanted state of mind or unwanted feeling.",
            ], "body"),
            ("", [
                "Think about the habit, state of mind, or feeling you want to change. Next, create an affirmation that is positive and opposite of the habit. ",
                "For example; to stop smoking, “I always take good care of my lungs.” To lose weight, “I only eat when I am hungry,” or, “I eat only healthy foods.”",
                "Think of the Seiheki symbol or draw it in the air or on both hands. Then place your dominant hand on the back of your head and the other hand on your forehead. ",
                "State the affirmation out loud or to yourself confidently several times. Then remove the non-dominant hand from the forehead and continue giving yourself Reiki to the back of the head for several minutes or as long as you are guided to. ",
            ], "number"),
        ],
    },
    {
        "kind": "symbol",
        "title": "THE DISTANT HEALING SYMBOL",
        "subtitle": "HON SHA ZE SHO NEN",
        "quote": "(HUN SHA ZAY SHOW NEN)<br/>“The Divine in me reaches out to the divine in you<br/>to promote enlightenment and peace.”",
        "blocks": [
            ("", [
                "The distant healing, or absentee symbol is used to send Reiki to others at a distance. It transcends time and space. This symbol taps into the levels which exist beyond time and space; the levels of the universal mind. It is a bridge to those at a distance, the past, the present, or the future. ",
            ], "body"),
            ("Ways to send distant healing;", [
                "Draw the symbol in the air or over a picture or object. Then beam Reiki with your hands towards the picture or object.",
                "Draw the symbol over the picture or object; hold it between your hands and request that the energy flow to the person or situation.",
                "Write the person’s name or describe the situation on a piece of paper, draw the distant healing symbol over it and hold the paper between your hands. ",
                "Hold your hands in the direction in which you imagine the person to be, draw the distant healing symbol and beam Reiki to them.",
            ], "bullet"),
            ("Manifesting Goals;", [
                "Simply write out your goal on a piece of paper and draw the symbol over it. Then the other 2 symbols. Then give the situation over to Reiki and give Reiki. Feel the energy flow. ",
                "If you know the date and time of a future event such as an interview, doctor’s appointment, testing or surgery, you can use the distant healing symbol to send energy that will be stored up like a battery for that important event. ",
                "Other Uses: sending Reiki to a client before they come. To attract clients. Send to traumatic events in the past. To higher dimensions. To someone who has passed on. To prepare a room for a lecture or session. ",
            ], "bullet"),
            ("", [
                "When you send Reiki to a traumatic event in the past, it helps heal the emotional and mental damage that remains today. Reiki helps offer new choices and actions. It helps to heal from physical trauma. It can help you heal karma. ",
                "If you are blocked in some area of your life you can draw the distant healing symbol over yourself and say.. “I now send Reiki to that part of myself that needs to heal in order for me to achieve ________.”",
            ], "body"),
        ],
    },
    {
        "kind": "practice",
        "title": "DISTANT HEALING",
        "subtitle": "ENKAKU CHIRYO",
        "quote": "",
        "blocks": [
            ("When sending Reiki to a person;", [
                "Do the Gasho (clear, center, focus) & Reijo Ho (inviting energy and intuition in). ",
                "Draw or visualize all of the symbols in your hands.",
                "Begin by drawing or visualizing the Distant Reiki symbol. Say it’s name 3x. Ask it to act like a bridge, connecting you to the person. Then connect with their higher self - visualize them as best as you can and ask permission to send them Reiki (this is not necessary if they have personally asked you already). ",
                "Aim your hands out beaming towards the person as you visualize them. ",
                "Send Reiki",
                "You can imagine sending each of their chakras Reiki, or scan their body and see where they need it and beam the Reiki there.",
                "Or you can send Reiki for 5-10 mins to each layer of their auric field. The physical, mental, emotional and spiritual. (PEMS)",
                "You may add other symbols if and when you feel guided. ",
                "When you are finished, seal the session with the power symbol, Chokurei, saying “I seal this healing with love and light” 3x. ",
                "Do kenyoku or dry bathing to clear and disconnect.",
            ], "number"),
        ],
    },
    {
        "kind": "session",
        "title": "Conducting a Full Reiki Session",
        "subtitle": "",
        "quote": "",
        "blocks": [
            ("Prepare yourself!", [
                "Put the power symbol or all of the symbols in both of the palms and gently tap or clap it in 3x. ",
                "Draw a large Chokurei over the body with the spiral ending at the solar plexus. ",
                "Draw a small Chokurei over each chakra- starting at the root and ending at the crown.",
                "Check in with your body- Are you grounded? Are you shielded/protected?",
            ], "bullet"),
            ("Prepare the room", [
                "Draw the Chokurei on each wall, the ceiling and the floor. ",
                "Draw all 3 symbols in the middle of the room. Feel their energies fill the room and clear out any negative energies. ",
            ], "bullet"),
            ("Set intention", [
                "Gasho- hands together at heart, igniting palm chakras, asking for your ego to step aside and setting the intention that the healing be for the highest good of the receiver.",
                "Reiji Ho- prayer hands at third eye- indication of spirit. Asking that you are intuitively guided to go where the energy needs to flow. Calling on any guides that you reside with to assist in the process (if this resonates with you). ",
            ], "bullet"),
            ("Conduct the session", [
                "Option to scan the person first, or just to start the process of a full body treatment using the designated hand position or places you feel guided. ",
            ], "bullet"),
            ("End the session ", [
                "Put one hand on the low back and another on the back of the heart. Draw or visualize the Chokurei encompassing the back and say in your mind 3x, “I seal this healing with love and light.”",
                "Gently take your hands off the person and allow them time to refocus. ",
                "Brush or sweep the energy from around the person and then also yourself (Kenyoku).",
            ], "bullet"),
        ],
    },
]


def html_escape(text):
    return (
        text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace("\n", "<br/>")
    )


def set_run(run, size=None, bold=None, color=None, font=None, spacing=None):
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = color
    if font:
        run.font.name = font
        run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    if spacing:
        rpr = run._r.get_or_add_rPr()
        sp = OxmlElement("w:spacing")
        sp.set(qn("w:val"), str(spacing))
        rpr.append(sp)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=BORDER, size="8"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_cell_margins(cell, top=120, start=180, bottom=120, end=180):
    tc_pr = cell._tc.get_or_add_tcPr()
    mar = tc_pr.first_child_found_in("w:tcMar")
    if mar is None:
        mar = OxmlElement("w:tcMar")
        tc_pr.append(mar)
    for name, val in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            mar.append(node)
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, width):
    tbl_w = table._tbl.tblPr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        table._tbl.tblPr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width))
    tbl_w.set(qn("w:type"), "dxa")


def add_docx_paragraph(parent, text="", size=10.5, color=MUTED, font="Avenir Next", after=6, align=None):
    p = parent.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.12
    if text:
        run = p.add_run(text)
        set_run(run, size=size, color=color, font=font)
    return p


def docx_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, 9000)
    cell = table.cell(0, 0)
    shade_cell(cell, SOFT)
    set_cell_border(cell)
    set_cell_margins(cell, top=160, bottom=160, start=220, end=220)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(title)
    set_run(r, size=9, bold=True, color=TEXT, font="Avenir Next", spacing=80)
    add_docx_paragraph(cell, body, after=0)


def build_docx():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(0.65)
    sec.bottom_margin = Inches(0.65)
    sec.left_margin = Inches(0.72)
    sec.right_margin = Inches(0.72)
    doc.styles["Normal"].font.name = "Avenir Next"
    doc.styles["Normal"].font.size = Pt(10.5)

    header = sec.header.paragraphs[0]
    r = header.add_run("T H E  L I G H T N E S S  O F  B E I N G")
    set_run(r, size=8, bold=True, color=MUTED, font="Avenir Next")

    add_docx_paragraph(doc, "", after=24)
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, 9000)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade_cell(cell, "FFFFFF")
    set_cell_border(cell)
    set_cell_margins(cell, top=780, bottom=620, start=360, end=360)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Usui/Holy Fire III")
    set_run(r, size=31, color=TEXT, font="Georgia")
    p = add_docx_paragraph(cell, "Reiki 2 Training Outline", size=18, color=TEXT, font="Georgia", align=WD_ALIGN_PARAGRAPH.CENTER, after=70)
    add_docx_paragraph(cell, "Led by Kate Gajewski", align=WD_ALIGN_PARAGRAPH.CENTER, after=130)
    add_docx_paragraph(cell, "bethelightness.com", align=WD_ALIGN_PARAGRAPH.CENTER, after=3)
    add_docx_paragraph(cell, "2026", align=WD_ALIGN_PARAGRAPH.CENTER, after=0)

    doc.add_page_break()
    p = add_docx_paragraph(doc, "Training Outline", size=20, color=TEXT, font="Georgia", after=10)
    for item in outline:
        p = add_docx_paragraph(doc, item, size=10.5, color=TEXT, after=4)
        p.style = "List Number"

    for section in sections:
        doc.add_page_break()
        add_docx_paragraph(doc, "REIKI SYMBOL" if section["kind"] == "symbol" else "PRACTICE", size=8, color=MUTED, font="Avenir Next", after=4)
        add_docx_paragraph(doc, section["title"], size=21, color=TEXT, font="Georgia", after=5)
        if section["subtitle"]:
            docx_callout(doc, section["subtitle"], section["quote"].replace("<br/>", "\n"))
        elif section["quote"]:
            docx_callout(doc, "", section["quote"].replace("<br/>", "\n"))
        for heading, items, style in section["blocks"]:
            if heading:
                add_docx_paragraph(doc, heading, size=14, color=TEXT, font="Georgia", after=5)
            for idx, item in enumerate(items, 1):
                p = add_docx_paragraph(doc, item, size=10.5, color=MUTED, after=5)
                if style == "bullet":
                    p.style = "List Bullet"
                elif style == "number":
                    p.style = "List Number"
    add_docx_paragraph(doc, "Kate Gajewski, 2024", align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.save(DOCX_OUT)


pdf_styles = getSampleStyleSheet()
pdf_styles.add(ParagraphStyle("H1x", fontName="Times-Bold", fontSize=21, leading=25, textColor=colors.HexColor(TEXT_HEX), spaceAfter=8))
pdf_styles.add(ParagraphStyle("H2x", fontName="Times-Bold", fontSize=14, leading=17, textColor=colors.HexColor(TEXT_HEX), spaceBefore=10, spaceAfter=5))
pdf_styles.add(ParagraphStyle("Bodyx", fontName="Helvetica", fontSize=10.5, leading=14.3, textColor=colors.HexColor(MUTED_HEX), spaceAfter=6))
pdf_styles.add(ParagraphStyle("Kickerx", fontName="Helvetica-Bold", fontSize=8, leading=10, textColor=colors.HexColor(MUTED_HEX), spaceBefore=4, spaceAfter=4))
pdf_styles.add(ParagraphStyle("Listx", parent=pdf_styles["Bodyx"], leftIndent=18, firstLineIndent=-10, spaceAfter=5))
pdf_styles.add(ParagraphStyle("Centerx", parent=pdf_styles["Bodyx"], alignment=1))


def pdf_p(text, style="Bodyx"):
    return Paragraph(html_escape(text), pdf_styles[style])


def pdf_bullet(text):
    return Paragraph("• " + html_escape(text), pdf_styles["Listx"])


def pdf_number(text, n):
    return Paragraph(f"{n}. " + html_escape(text), pdf_styles["Listx"])


def draw_header(canvas, doc):
    canvas.saveState()
    canvas.setFont("Helvetica-Bold", 7.5)
    canvas.setFillColor(colors.HexColor(MUTED_HEX))
    canvas.drawString(0.72 * inch, letter[1] - 0.47 * inch, "T H E  L I G H T N E S S  O F  B E I N G")
    canvas.setFont("Helvetica", 8)
    canvas.drawRightString(letter[0] - 0.72 * inch, 0.38 * inch, str(doc.page))
    canvas.restoreState()


def draw_title(canvas, doc):
    draw_header(canvas, doc)
    canvas.saveState()
    x, y = 0.75 * inch, 0.92 * inch
    w, h = letter[0] - 1.5 * inch, letter[1] - 1.65 * inch
    canvas.setStrokeColor(colors.HexColor(BORDER_HEX))
    canvas.roundRect(x, y, w, h, 14, stroke=1, fill=0)
    canvas.setFillColor(colors.HexColor(TEXT_HEX))
    canvas.setFont("Times-Roman", 31)
    canvas.drawCentredString(letter[0] / 2, letter[1] * 0.66, "Usui/Holy Fire III")
    canvas.setFont("Times-Roman", 18)
    canvas.drawCentredString(letter[0] / 2, letter[1] * 0.61, "Reiki 2 Training Outline")
    canvas.setFillColor(colors.HexColor(MUTED_HEX))
    canvas.setFont("Helvetica", 10)
    canvas.drawCentredString(letter[0] / 2, letter[1] * 0.52, "Led by Kate Gajewski")
    canvas.drawCentredString(letter[0] / 2, letter[1] * 0.16, "bethelightness.com")
    canvas.drawCentredString(letter[0] / 2, letter[1] * 0.135, "2026")
    canvas.restoreState()


def pdf_callout(title, body):
    table = Table(
        [[pdf_p(title.upper(), "Kickerx"), pdf_p(body.replace("<br/>", "\n"), "Bodyx")]],
        colWidths=[1.55 * inch, 5.05 * inch],
        hAlign="LEFT",
    )
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor(SOFT_HEX)),
        ("BOX", (0, 0), (-1, -1), 0.6, colors.HexColor(BORDER_HEX)),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("LEFTPADDING", (0, 0), (-1, -1), 13),
        ("RIGHTPADDING", (0, 0), (-1, -1), 13),
        ("TOPPADDING", (0, 0), (-1, -1), 12),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 12),
    ]))
    return [table, Spacer(1, 8)]


def build_pdf():
    doc = BaseDocTemplate(PDF_OUT, pagesize=letter, leftMargin=0.72 * inch, rightMargin=0.72 * inch, topMargin=0.72 * inch, bottomMargin=0.7 * inch)
    frame = Frame(doc.leftMargin, doc.bottomMargin, doc.width, doc.height - 0.2 * inch, id="normal")
    doc.addPageTemplates([PageTemplate(id="title", frames=frame, onPage=draw_title), PageTemplate(id="body", frames=frame, onPage=draw_header)])
    story = [NextPageTemplate("body"), PageBreak(), pdf_p("Training Outline", "H1x")]
    for i, item in enumerate(outline, 1):
        story.append(pdf_number(item, i))
    for section in sections:
        story.append(PageBreak())
        story.append(pdf_p("REIKI SYMBOL" if section["kind"] == "symbol" else "PRACTICE", "Kickerx"))
        story.append(pdf_p(section["title"], "H1x"))
        if section["subtitle"] or section["quote"]:
            story += pdf_callout(section["subtitle"], section["quote"])
        for heading, items, style in section["blocks"]:
            if heading:
                story.append(pdf_p(heading, "H2x"))
            for idx, item in enumerate(items, 1):
                if style == "bullet":
                    story.append(pdf_bullet(item))
                elif style == "number":
                    story.append(pdf_number(item, idx))
                else:
                    story.append(pdf_p(item))
    story += [Spacer(1, 18), pdf_p("Kate Gajewski, 2024", "Centerx")]
    doc.build(story)


if __name__ == "__main__":
    build_docx()
    build_pdf()
    print(DOCX_OUT)
    print(PDF_OUT)
